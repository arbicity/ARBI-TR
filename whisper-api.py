from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse
import shutil
import subprocess
import tempfile
import datetime
import os
import time
from typing import Tuple, List, Optional
import numpy as np
from docx import Document
import pandas as pd
import torch
import torchaudio
from sklearn.cluster import AgglomerativeClustering
from transformers import pipeline
from pyannote.audio.pipelines.speaker_verification import PretrainedSpeakerEmbedding
import multiprocessing
from functools import partial

app = FastAPI()

def convert_time(secs: float) -> str:
    if secs is None:
        return "00:00:00"  # Handle None values gracefully
    return str(datetime.timedelta(seconds=round(secs)))

def timeit(func):
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        print(f"{func.__name__} executed in {time.time() - start_time:.4f} seconds")
        return result
    return wrapper

@timeit
def convert_audio_to_wav(media_file_path: str) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_wav_file:
        audio_file_path = temp_wav_file.name
    try:
        cmd = ['ffmpeg', '-y', '-i', media_file_path, '-ar', '16000', '-ac', '1', '-c:a', 'pcm_s16le', '-c:v', 'hevc_nvenc', '-preset', 'fast', audio_file_path]
        subprocess.run(cmd, check=True)
        print(f"Converted {media_file_path} to WAV: {audio_file_path}")
    except Exception as e:
        os.remove(audio_file_path)
        raise e
    return audio_file_path

@timeit
def load_audio_file(file_path: str) -> Tuple[torch.Tensor, int]:
    return torchaudio.load(file_path)

@timeit
def perform_transcription(asr_pipeline, converted_audio_file_path: str):
    """Performs transcription using the specified ASR pipeline."""
    return asr_pipeline(converted_audio_file_path, chunk_length_s=30, batch_size=48, return_timestamps=True)

@timeit
def generate_transcription_output(segments, speaker_labels) -> dict:
    """Generates transcription output with speaker labels, reorganizing for grouped output."""
    output_segments = []
    text, current_speaker, prev_start_time, prev_end_time = '', None, '', ''
    for i, segment in enumerate(segments):
        start_time = convert_time(segment["timestamp"][0])
        end_time = convert_time(segment["timestamp"][1])
        speaker = speaker_labels[i]

        if current_speaker is None or current_speaker != speaker:
            if text:  # If there's accumulated text, append the previous segment before starting a new one
                output_segments.append({
                    "Start": prev_start_time,
                    "End": prev_end_time,
                    "Speaker": str(current_speaker),
                    "Text": text.strip()
                })
                text = ''  # Reset text for the new segment

            current_speaker = speaker  # Update the current speaker
            prev_start_time = start_time  # Update start time for potential new segment

        text += segment["text"] + ' '
        prev_end_time = end_time  # Always update end time to the latest for ongoing speaker segments

    # After loop, add the last accumulated segment
    if text:
        output_segments.append({
            "Start": prev_start_time,
            "End": prev_end_time,
            "Speaker": str(current_speaker),
            "Text": text.strip()
        })

    return {"segments": output_segments}

@timeit
def generate_embeddings_for_segments(embedding_model, waveform: torch.Tensor, sample_rate: int, segments: List[Tuple[float, Optional[float]]]) -> List[torch.Tensor]:
    """Generates embeddings for a list of given audio segments."""
    embeddings = []
    audio_length_in_seconds = waveform.shape[1] / sample_rate  # Calculate total audio duration

    for start_time, end_time in segments:
        if end_time is None or end_time > audio_length_in_seconds:
            end_time = audio_length_in_seconds  # Default to the end of the audio if end_time is None or out of bounds

        start_sample = int(start_time * sample_rate)
        end_sample = int(end_time * sample_rate)
        segment_waveform = waveform[:, start_sample:end_sample].unsqueeze(0)
        embedding = embedding_model(segment_waveform)
        embeddings.append(embedding)
    return embeddings

def cluster_embeddings(embeddings: List[torch.Tensor]) -> np.ndarray:
    """Clusters embeddings and returns the labels."""
    embeddings_array = np.vstack(embeddings)  # Directly use embeddings if they're numpy arrays
    clustering_model = AgglomerativeClustering(n_clusters=None, compute_full_tree=True, distance_threshold=2000)
    clustering_model.fit(embeddings_array)
    print(f"Clustering completed. Labels: {np.unique(clustering_model.labels_)}")
    return clustering_model.labels_

def process_audio(file_path: str, size_of_model: str, task: str, source_language: Optional[str], speaker_number: Optional[int]) -> dict:
    converted_audio_file_path = convert_audio_to_wav(file_path)

    model_id = f"openai/whisper-{size_of_model}" + ("-v3" if size_of_model == "large" else "")

    asr_pipeline = pipeline(
        "automatic-speech-recognition",
        model=model_id,
        torch_dtype=torch.float16,
        model_kwargs={
            "device_map": "cuda:0" if torch.cuda.is_available() else "cpu",
            "attn_implementation": "sdpa"
        },
        generate_kwargs={
            "task": task,
            "language": source_language if source_language else None
        }
    )

    waveform, sample_rate = load_audio_file(converted_audio_file_path)
    
    transcription_result = perform_transcription(asr_pipeline, converted_audio_file_path)

    valid_chunks = [chunk for chunk in transcription_result["chunks"] if chunk["timestamp"][1] is not None]
    segments = [(chunk["timestamp"][0], chunk["timestamp"][1]) for chunk in valid_chunks]

    if speaker_number == 0:
        embedding_model = PretrainedSpeakerEmbedding("speechbrain/spkrec-ecapa-voxceleb", device="cuda")
    else:
        embedding_model = PretrainedSpeakerEmbedding("speechbrain/spkrec-ecapa-voxceleb", device="cuda")
    
    embeddings = generate_embeddings_for_segments(embedding_model, waveform, sample_rate, segments)
    
    if speaker_number == 0:
        speaker_labels = cluster_embeddings(embeddings)
    else:
        clustering_model = AgglomerativeClustering(n_clusters=speaker_number)
        speaker_labels = clustering_model.fit_predict(np.vstack(embeddings))

    grouped_segments = generate_transcription_output(valid_chunks, speaker_labels)
    
    del embedding_model
    
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
    
    os.remove(file_path)
    os.remove(converted_audio_file_path)
    
    return grouped_segments

def create_transcript_doc(df, filename="transcription.docx"):
    """
    Create a Word document from a DataFrame.
    df: DataFrame with columns 'Start', 'End', 'Speaker', and 'Text'.
    """
    document = Document()
    document.add_heading('Transcription', level=0)

    # Add a table to the document
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Start'
    hdr_cells[1].text = 'End'
    hdr_cells[2].text = 'Speaker'
    hdr_cells[3].text = 'Text'

    # Populate the table with rows from the DataFrame
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Start'])
        row_cells[1].text = str(row['End'])
        row_cells[2].text = str(row['Speaker'])
        row_cells[3].text = row['Text']

    document.save(filename)
    return filename
@app.post("/transcribe/")
async def transcribe_audio(file: UploadFile = File(...), 
                           size_of_model: str = Form(...), 
                           task: str = Form(...), 
                           source_language: Optional[str] = Form(None), 
                           speaker_number: Optional[int] = Form(0)):
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        shutil.copyfileobj(file.file, temp_file)
        temp_file_path = temp_file.name

    with multiprocessing.Pool(1) as pool:
        async_result = pool.apply_async(process_audio, (temp_file_path, size_of_model, task, source_language, speaker_number))
        result = async_result.get()  # Wait for the process to complete and get the result

    # Convert the result to a pandas DataFrame for easier manipulation
    df = pd.DataFrame(result['segments'])
    # Generate a Word document from the DataFrame
    doc_path = create_transcript_doc(df)
    
    # Return the Word document as a downloadable file
    return FileResponse(path=doc_path, filename="transcription.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
